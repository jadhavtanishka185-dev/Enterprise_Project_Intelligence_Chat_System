# import os
# from pathlib import Path
# from typing import List
# from langchain_core.documents import Document
# from langchain_text_splitters import RecursiveCharacterTextSplitter


# def extract_text_from_pdf(filepath: str) -> List[Document]:
#     from pypdf import PdfReader
#     reader = PdfReader(filepath)
#     docs = []
#     for page_num, page in enumerate(reader.pages):
#         text = page.extract_text() or ""
#         if text.strip():
#             docs.append(Document(
#                 page_content=text,
#                 metadata={"source": Path(filepath).name, "page": page_num + 1}
#             ))
#     return docs


# def extract_text_from_docx(filepath: str) -> List[Document]:
#     from docx import Document as DocxDocument
#     doc = DocxDocument(filepath)
#     full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
#     return [Document(
#         page_content=full_text,
#         metadata={"source": Path(filepath).name, "page": 1}
#     )]


# def extract_text_from_txt(filepath: str) -> List[Document]:
#     with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
#         text = f.read()
#     return [Document(
#         page_content=text,
#         metadata={"source": Path(filepath).name, "page": 1}
#     )]


# def load_document(filepath: str) -> List[Document]:
#     ext = Path(filepath).suffix.lower()
#     if ext == ".pdf":
#         return extract_text_from_pdf(filepath)
#     elif ext in (".docx", ".doc"):
#         return extract_text_from_docx(filepath)
#     elif ext == ".txt":
#         return extract_text_from_txt(filepath)
#     else:
#         raise ValueError(f"Unsupported file type: {ext}")


# def split_documents(docs: List[Document], chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Document]:
#     splitter = RecursiveCharacterTextSplitter(
#         chunk_size=chunk_size,
#         chunk_overlap=chunk_overlap,
#         separators=["\n\n", "\n", ". ", " ", ""],
#     )
#     return splitter.split_documents(docs)



import os
from pathlib import Path
from typing import List

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


# =========================================================
# PDF
# =========================================================

def extract_text_from_pdf(
    filepath: str
) -> List[Document]:

    from pypdf import PdfReader

    reader = PdfReader(filepath)

    docs = []

    for page_num, page in enumerate(reader.pages):

        text = page.extract_text() or ""

        if text.strip():

            docs.append(
                Document(
                    page_content=text,

                    metadata={
                        "source": str(
                            Path(filepath).name
                        ),

                        # IMPORTANT:
                        # Convert to string
                        "page": str(
                            page_num + 1
                        ),
                    }
                )
            )

    return docs


# =========================================================
# DOCX
# =========================================================

def extract_text_from_docx(
    filepath: str
) -> List[Document]:

    from docx import Document as DocxDocument

    doc = DocxDocument(filepath)

    full_text = "\n".join([
        para.text
        for para in doc.paragraphs
        if para.text.strip()
    ])

    return [
        Document(
            page_content=full_text,

            metadata={
                "source": str(
                    Path(filepath).name
                ),

                # IMPORTANT:
                # Convert to string
                "page": "1",
            }
        )
    ]


# =========================================================
# TXT
# =========================================================

def extract_text_from_txt(
    filepath: str
) -> List[Document]:

    with open(
        filepath,
        "r",
        encoding="utf-8",
        errors="ignore"
    ) as f:

        text = f.read()

    return [
        Document(
            page_content=text,

            metadata={
                "source": str(
                    Path(filepath).name
                ),

                # IMPORTANT:
                # Convert to string
                "page": "1",
            }
        )
    ]


# =========================================================
# LOAD DOCUMENT
# =========================================================

def load_document(
    filepath: str
) -> List[Document]:

    ext = Path(filepath).suffix.lower()

    if ext == ".pdf":

        return extract_text_from_pdf(
            filepath
        )

    elif ext in (
        ".docx",
        ".doc"
    ):

        return extract_text_from_docx(
            filepath
        )

    elif ext == ".txt":

        return extract_text_from_txt(
            filepath
        )

    else:

        raise ValueError(
            f"Unsupported file type: {ext}"
        )


# =========================================================
# SPLIT DOCUMENTS
# =========================================================

def split_documents(
    docs: List[Document],
    chunk_size: int = 1000,
    chunk_overlap: int = 200
) -> List[Document]:

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=[
            "\n\n",
            "\n",
            ". ",
            " ",
            ""
        ],
    )

    split_docs = splitter.split_documents(
        docs
    )

    # Remove empty chunks
    split_docs = [
        doc
        for doc in split_docs
        if doc.page_content.strip()
    ]

    return split_docs